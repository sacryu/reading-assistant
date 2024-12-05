from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from docx import Document
import io

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 允许所有来源
    allow_credentials=True,
    allow_methods=["*"],  # 允许所有HTTP方法
    allow_headers=["*"],  # 允许所有HTTP头
)


@app.post("/convert")
async def convert(file: UploadFile = File(...)):
    if file.content_type == "application/pdf":
        return await convert_pdf_to_html(file)
    elif file.content_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]:
        return await convert_docx_to_html(file)
    else:
        return JSONResponse(
            status_code=400, content={"message": "Unsupported file type"}
        )

async def convert_pdf_to_html(file: UploadFile):
    pass

async def convert_docx_to_html(file: UploadFile):
    file_bytes = await file.read()
    file_stream = io.BytesIO(file_bytes)
    document = Document(file_stream)
    
    html = "<html><body>"
    for para in document.paragraphs:
        html += f"<p>{para.text}</p>"
    for table in document.tables:
        html += "<table border='1'>"
        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                html += f"<td>{cell.text}</td>"
            html += "</tr>"
        html += "</table>"
    html += "</body></html>"

    return {"html": html}

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app="main:app", host="0.0.0.0", port=8000, reload=True)